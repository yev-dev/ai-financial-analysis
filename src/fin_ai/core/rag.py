import os
import json
import logging
import shutil
import tempfile
import warnings
import json
import xml.etree.ElementTree as ET
from pathlib import Path
from datetime import datetime

os.environ.setdefault("KMP_DUPLICATE_LIB_OK", "TRUE")

warnings.filterwarnings(
    "ignore",
    message=r"Accessing `__path__` from `\.models\..*",
)

from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_text_splitters import MarkdownHeaderTextSplitter
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.documents import Document
from langchain_ollama import ChatOllama
import faiss
import pymupdf
import pandas as pd
from bs4 import BeautifulSoup
import html2text
from docx import Document as WordDocument

from dashboard import DEFAULT_CHAT_MODEL, OLLAMA_BASE_URL, VECTOR_DB_DIR

def _embedding_metadata_path(filename: str) -> Path:
    return Path(VECTOR_DB_DIR) / f"{filename}.embedding.json"


def save_embedding_metadata(filename: str, provider: str, model: str, base_url: str) -> Path:
    metadata_path = _embedding_metadata_path(filename)
    metadata = {
        "provider": provider,
        "model": model,
        "base_url": base_url.rstrip("/"),
    }
    metadata_path.write_text(json.dumps(metadata, ensure_ascii=True, indent=2), encoding="utf-8")
    return metadata_path


def load_embedding_metadata(filename: str) -> dict[str, str] | None:
    metadata_path = _embedding_metadata_path(filename)
    if not metadata_path.exists():
        return None

    try:
        raw = json.loads(metadata_path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return None

    provider = str(raw.get("provider", "")).strip()
    model = str(raw.get("model", "")).strip()
    base_url = str(raw.get("base_url", "")).strip().rstrip("/")
    if not provider or not model or not base_url:
        return None

    return {
        "provider": provider,
        "model": model,
        "base_url": base_url,
    }


def _load_and_convert_with_pymupdf(file_path):
    document = pymupdf.open(file_path)
    sections = []
    try:
        for page_number, page in enumerate(document, start=1):
            text = page.get_text().strip()
            sections.append(f"# Page {page_number}\n\n{text}")
    finally:
        document.close()

    return "\n\n".join(section for section in sections if section.strip())


# Load and convert PDF, CSV, JSON, XML, HTML, Word or URL content to markdown
def load_and_convert_document(file_path):
    """Load and convert a document to markdown format.
    
    Supports:
    - PDF files (.pdf)
    - CSV files (.csv)
    - JSON files (.json)
    - XML files (.xml)
    - HTML files (.html)
    - Word files (.docx)
    - URLs (http://, https://)
    
    Args:
        file_path: Path to file or URL
        
    Returns:
        Markdown formatted content
        
    Raises:
        ValueError: If file format is not supported
    """
    # Check if it's a URL
    if isinstance(file_path, str) and (file_path.startswith('http://') or file_path.startswith('https://')):
        html_content = _scrape_html_from_url(file_path)
        return _convert_html_to_markdown(html_content, source_url=file_path)
    
    source_path = Path(file_path)
    file_extension = source_path.suffix.lower()
    
    if file_extension == ".csv":
        return _convert_csv_to_markdown(source_path)
    elif file_extension == ".json":
        return _convert_json_to_markdown(source_path)
    elif file_extension == ".xml":
        return _convert_xml_to_markdown(source_path)
    elif file_extension == ".docx":
        return _convert_word_to_markdown(source_path)
    elif file_extension == ".pdf":
        with tempfile.TemporaryDirectory(prefix="fin_ai_pdf_") as temp_dir:
            temp_pdf_path = Path(temp_dir) / source_path.name
            shutil.copy2(source_path, temp_pdf_path)
            return _convert_pdf_to_markdown(temp_pdf_path)
    elif file_extension == ".html":
        with open(source_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        return _convert_html_to_markdown(html_content, source_url=str(source_path))
    else:
        raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: .pdf, .csv, .json, .xml, .html, .docx or URLs")


def _convert_pdf_to_markdown(file_path: Path) -> str:
    try:
        from docling.document_converter import DocumentConverter

        converter = DocumentConverter()
        result = converter.convert(file_path)
        return result.document.export_to_markdown()
    except Exception as exc:
        warnings.warn(
            f"Falling back to PyMuPDF text extraction because docling conversion failed: {exc}",
            RuntimeWarning,
            stacklevel=2,
        )
        return _load_and_convert_with_pymupdf(file_path)


def _convert_csv_to_markdown(file_path: Path) -> str:
    """Convert CSV file to markdown format.
    
    Args:
        file_path: Path to the CSV file
        
    Returns:
        Markdown formatted string representation of the CSV data
    """
    try:
        df = pd.read_csv(file_path)
        
        # Create markdown header with filename and metadata
        filename = file_path.name
        markdown_content = f"# {filename}\n\n"
        markdown_content += f"## Data Overview\n"
        markdown_content += f"- Rows: {len(df)}\n"
        markdown_content += f"- Columns: {len(df.columns)}\n"
        markdown_content += f"- Column Names: {', '.join(df.columns)}\n\n"
        
        # Add column data types
        markdown_content += "## Column Information\n"
        for col in df.columns:
            markdown_content += f"- {col}: {df[col].dtype}\n"
        markdown_content += "\n"
        
        # Add data as markdown table with chunking for large datasets
        markdown_content += "## Data\n"
        if len(df) > 100:
            # For large datasets, add summary statistics
            markdown_content += "### Summary Statistics\n"
            markdown_content += df.describe().to_markdown()
            markdown_content += "\n\n### Sample Data (First 50 Rows)\n"
            markdown_content += df.head(50).to_markdown(index=False)
        else:
            markdown_content += df.to_markdown(index=False)
        
        return markdown_content
    except Exception as exc:
        raise ValueError(f"Failed to convert CSV file {file_path}: {exc}") from exc


def _json_to_markdown_recursive(obj, level=1, max_depth=5):
    """Recursively convert JSON object to markdown format.
    
    Args:
        obj: JSON object (dict, list, or primitive)
        level: Current markdown header level
        max_depth: Maximum recursion depth to prevent excessive nesting
        
    Returns:
        Markdown formatted string
    """
    if level > max_depth:
        return f"`[{type(obj).__name__}]`"
    
    markdown = ""
    
    if isinstance(obj, dict):
        for key, value in obj.items():
            # Create header for dict keys
            header = "#" * (level + 1)
            markdown += f"{header} {key}\n\n"
            
            if isinstance(value, (dict, list)):
                markdown += _json_to_markdown_recursive(value, level + 1, max_depth)
            else:
                markdown += f"{value}\n\n"
    
    elif isinstance(obj, list):
        if not obj:
            markdown += "*(empty list)*\n\n"
        else:
            # For lists of primitives, show as bullet points
            if all(isinstance(item, (str, int, float, bool, type(None))) for item in obj):
                for item in obj:
                    markdown += f"- {item}\n"
                markdown += "\n"
            else:
                # For lists of objects, process each item
                for idx, item in enumerate(obj[:50]):  # Limit to first 50 items
                    markdown += f"{chr(96)}Item {idx + 1}{chr(96)}\n"
                    if isinstance(item, (dict, list)):
                        markdown += _json_to_markdown_recursive(item, level + 1, max_depth)
                    else:
                        markdown += f"{item}\n\n"
                
                if len(obj) > 50:
                    markdown += f"*... and {len(obj) - 50} more items*\n\n"
    
    else:
        markdown += f"{obj}\n\n"
    
    return markdown


def _convert_json_to_markdown(file_path: Path) -> str:
    """Convert JSON file to markdown format.
    
    Args:
        file_path: Path to the JSON file
        
    Returns:
        Markdown formatted string representation of the JSON data
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Create markdown header with filename
        filename = file_path.name
        markdown_content = f"# {filename}\n\n"
        
        # Get file statistics
        file_size = file_path.stat().st_size
        markdown_content += f"## File Information\n"
        markdown_content += f"- Size: {file_size:,} bytes\n"
        markdown_content += f"- Type: {type(data).__name__}\n\n"
        
        # Convert JSON content to markdown
        markdown_content += "## Content\n\n"
        markdown_content += _json_to_markdown_recursive(data, level=1)
        
        return markdown_content
    except json.JSONDecodeError as exc:
        raise ValueError(f"Invalid JSON file {file_path}: {exc}") from exc
    except Exception as exc:
        raise ValueError(f"Failed to convert JSON file {file_path}: {exc}") from exc


def _xml_to_markdown_recursive(element: ET.Element, level: int = 2, max_depth: int = 7) -> str:
    """Recursively convert an XML element tree to markdown."""
    if level > max_depth:
        return "`[Max depth reached]`\n\n"

    lines: list[str] = []

    heading = "#" * min(level, 6)
    lines.append(f"{heading} {element.tag}\n")

    if element.attrib:
        lines.append("**Attributes**\n")
        for key, value in element.attrib.items():
            lines.append(f"- {key}: {value}\n")
        lines.append("\n")

    text = (element.text or "").strip()
    if text:
        lines.append(f"{text}\n\n")

    children = list(element)
    if children:
        lines.append("**Children**\n\n")
        for child in children:
            lines.append(_xml_to_markdown_recursive(child, level + 1, max_depth))

    return "".join(lines)


def _convert_xml_to_markdown(file_path: Path) -> str:
    """Convert XML file to markdown format."""
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()

        file_size = file_path.stat().st_size
        markdown_content = f"# {file_path.name}\n\n"
        markdown_content += "## File Information\n"
        markdown_content += f"- Size: {file_size:,} bytes\n"
        markdown_content += f"- Root Tag: {root.tag}\n\n"
        markdown_content += "## Content\n\n"
        markdown_content += _xml_to_markdown_recursive(root)
        return markdown_content
    except ET.ParseError as exc:
        raise ValueError(f"Invalid XML file {file_path}: {exc}") from exc
    except Exception as exc:
        raise ValueError(f"Failed to convert XML file {file_path}: {exc}") from exc


def _convert_word_to_markdown(file_path: Path) -> str:
    """Convert DOCX file to markdown format."""
    try:
        document = WordDocument(file_path)
    except Exception as exc:
        raise ValueError(f"Failed to open Word file {file_path}: {exc}") from exc

    sections = [f"# {file_path.name}\n"]

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = getattr(paragraph.style, "name", "") or ""
        if style_name.startswith("Heading"):
            level_text = style_name.replace("Heading", "").strip()
            try:
                level = max(1, min(int(level_text), 6))
            except ValueError:
                level = 2
            sections.append(f"{'#' * level} {text}\n")
        else:
            sections.append(f"{text}\n")

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(values):
                rows.append(values)

        if not rows:
            continue

        sections.append(f"## Table {table_index}\n")
        header = rows[0]
        sections.append("| " + " | ".join(header) + " |\n")
        sections.append("| " + " | ".join("---" for _ in header) + " |\n")
        for row in rows[1:]:
            padded = row + [""] * (len(header) - len(row))
            sections.append("| " + " | ".join(padded[:len(header)]) + " |\n")

    markdown_content = "\n".join(section.strip("\n") for section in sections if section.strip())
    if not markdown_content.strip():
        raise ValueError(f"Word file {file_path} does not contain extractable content")
    return markdown_content + "\n"


def _scrape_html_from_url(url: str) -> str:
    """Scrape HTML content from a URL.
    
    Args:
        url: URL to scrape
        
    Returns:
        Raw HTML content
    """
    try:
        import requests
        
        headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        return response.text
    except Exception as exc:
        raise ValueError(f"Failed to scrape URL {url}: {exc}") from exc


def _convert_html_to_markdown(html_content: str, source_url: str = None) -> str:
    """Convert HTML content to markdown format.
    
    Args:
        html_content: HTML content to convert
        source_url: Optional URL source for reference
        
    Returns:
        Markdown formatted string
    """
    try:
        # Parse HTML and extract main content
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(['script', 'style']):
            script.decompose()
        
        # Remove comments
        for comment in soup.find_all(string=lambda text: isinstance(text, str) and text.strip().startswith('<!--')):
            comment.extract()
        
        # Extract title if available
        title = soup.find('title')
        title_text = title.get_text(strip=True) if title else "Web Content"
        
        # Get main content - try to find article/main tag, fallback to body
        main_content = soup.find(['article', 'main'])
        if not main_content:
            main_content = soup.find('body') or soup
        
        # Convert HTML to markdown using html2text
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = False
        h.body_width = 0  # Don't wrap text
        
        markdown_content = f"# {title_text}\n\n"
        
        if source_url:
            markdown_content += f"**Source**: {source_url}\n\n"
        
        # Convert main content to markdown
        html_str = str(main_content)
        converted = h.handle(html_str)
        markdown_content += converted
        
        return markdown_content
    except Exception as exc:
        raise ValueError(f"Failed to convert HTML to markdown: {exc}") from exc


def _create_source_metadata(file_path, source_type: str = "file"):
    """Create metadata for a document source.
    
    Args:
        file_path: Path to the source file or URL
        source_type: Type of source (file, url, csv, pdf, json, html)
        
    Returns:
        Dictionary with metadata
    """
    if isinstance(file_path, str) and (file_path.startswith('http://') or file_path.startswith('https://')):
        return {
            "source": file_path,
            "source_type": "url",
            "filename": file_path.split('/')[-1],
            "loaded_at": datetime.now().isoformat(),
        }
    
    path = Path(file_path)
    file_ext = path.suffix.lower()
    
    return {
        "source": str(path.absolute()),
        "source_type": source_type or file_ext[1:],
        "filename": path.name,
        "file_size": path.stat().st_size if path.exists() else 0,
        "loaded_at": datetime.now().isoformat(),
    }


# Split markdown into chunks with metadata
def get_markdown_splits(markdown_content, metadata: dict = None):
    """Split markdown content into chunks and add metadata to each chunk.
    
    Args:
        markdown_content: Markdown formatted string to split
        metadata: Optional metadata dictionary to attach to all chunks
        
    Returns:
        List of Document objects with metadata
    """
    headers_to_split_on = [("#", "Header 1"), ("##", "Header 2"), ("###", "Header 3")]
    splitter = MarkdownHeaderTextSplitter(headers_to_split_on, strip_headers=False)
    splits = splitter.split_text(markdown_content)
    
    # Convert to Document objects and add metadata
    if metadata is None:
        metadata = {}
    
    documents = []
    for idx, split in enumerate(splits):
        # Merge metadata with section metadata
        doc_metadata = {
            **metadata,
            "chunk_index": idx,
            "chunk_count": len(splits),
        }
        
        # Add header information if available
        if hasattr(split, 'metadata'):
            doc_metadata.update(split.metadata)
        
        doc = Document(
            page_content=split.page_content if hasattr(split, 'page_content') else split,
            metadata=doc_metadata
        )
        documents.append(doc)
    
    return documents


def _get_source_group_key(chunks: list) -> str:
    """Extract the source group key from document chunks.
    
    Uses the filename from the first chunk's metadata as the grouping key.
    This keeps all chunks from the same source file grouped together.
    
    Args:
        chunks: List of Document objects with metadata
        
    Returns:
        Source group key (e.g., "JAZZ-Equity-Research-Report", "NVDA_report")
    """
    if not chunks:
        return "unknown"
    
    for chunk in chunks:
        if isinstance(chunk, Document) and "filename" in chunk.metadata:
            filename = chunk.metadata["filename"]
            # Remove file extension to get the base name
            base_name = Path(filename).stem
            return base_name
    
    return "unknown"


def discover_vector_stores_by_source(vector_db_dir: str | Path = None) -> dict[str, Path]:
    """Discover all vector stores grouped by source in the vector database directory.
    
    Returns a dictionary mapping source names to their vector store paths.
    
    Args:
        vector_db_dir: Path to vector database directory (uses VECTOR_DB_DIR if None)
        
    Returns:
        Dictionary: {source_name: vector_store_path}
        Example: {"JAZZ-Equity-Research-Report": Path("vector_db/JAZZ-Equity-Research-Report/store.faiss"),
                  "NVDA_report": Path("vector_db/NVDA_report/store.faiss")}
    """
    if vector_db_dir is None:
        vector_db_dir = Path(VECTOR_DB_DIR)
    else:
        vector_db_dir = Path(vector_db_dir)
    
    vector_db_dir.mkdir(parents=True, exist_ok=True)
    
    sources = {}
    
    # Look for vector stores in source subdirectories (new structure)
    for source_dir in vector_db_dir.iterdir():
        if source_dir.is_dir() and not source_dir.name.startswith('.'):
            # Look for .faiss files in subdirectory
            faiss_files = list(source_dir.glob("*.faiss"))
            if faiss_files:
                # Use the first .faiss file found (or "store.faiss" by convention)
                store_path = next((f for f in faiss_files if f.name == "store.faiss"), faiss_files[0])
                sources[source_dir.name] = store_path
    
    # Legacy: also look for top-level .faiss directories (backward compatibility)
    for item in vector_db_dir.iterdir():
        if item.is_dir() and item.name.endswith('.faiss'):
            source_name = item.name.replace('.faiss', '')
            if source_name not in sources:
                sources[source_name] = item
    
    return sources


def get_or_create_source_vector_db_dir(source_group_key: str, vector_db_dir: str | Path = None) -> Path:
    """Get or create the vector database directory for a specific source.
    
    Args:
        source_group_key: Name of the source (e.g., "JAZZ-Equity-Research-Report")
        vector_db_dir: Parent vector database directory (uses VECTOR_DB_DIR if None)
        
    Returns:
        Path to the source's vector store directory
    """
    if vector_db_dir is None:
        vector_db_dir = Path(VECTOR_DB_DIR)
    else:
        vector_db_dir = Path(vector_db_dir)
    
    source_dir = vector_db_dir / source_group_key
    source_dir.mkdir(parents=True, exist_ok=True)
    return source_dir


def migrate_vector_stores_to_grouped_structure(vector_db_dir: str | Path = None) -> dict[str, str]:
    """Migrate existing vector stores from flat structure to source-grouped structure.
    
    This function reorganizes vector stores from:
        vector_db/{name}.faiss/ → vector_db/{name}/store.faiss
        vector_db/{name}.faiss → vector_db/{name}/store.faiss (new location)
    
    Args:
        vector_db_dir: Parent vector database directory (uses VECTOR_DB_DIR if None)
        
    Returns:
        Dictionary mapping source names to their new paths and migration status
    """
    if vector_db_dir is None:
        vector_db_dir = Path(VECTOR_DB_DIR)
    else:
        vector_db_dir = Path(vector_db_dir)
    
    migration_results = {}
    vector_db_dir.mkdir(parents=True, exist_ok=True)
    
    # Handle legacy .faiss directories (from FAISS.save_local())
    for item in vector_db_dir.iterdir():
        if item.is_dir() and item.name.endswith('.faiss'):
            source_name = item.name.replace('.faiss', '')
            target_dir = vector_db_dir / source_name
            
            # If target doesn't exist, rename the old directory
            if not target_dir.exists():
                try:
                    item.rename(target_dir)
                    migration_results[source_name] = f"Migrated {item.name} → {source_name}/"
                except OSError as e:
                    migration_results[source_name] = f"Error migrating {item.name}: {e}"
            else:
                migration_results[source_name] = f"Target {source_name}/ already exists, skipping"
    
    return migration_results



# Create or load the vector store
def create_or_load_vector_store(filename, chunks, embeddings, group_by_source: bool = True):
    """Create or load a vector store from chunks, optionally grouped by source.
    
    When group_by_source is True, vector stores are organized in source subdirectories
    (e.g., vector_db/JAZZ-Equity-Research-Report/store.faiss), keeping data grouped by
    source for better organization and retrieval boundaries.
    
    Args:
        filename: Name for the vector store
        chunks: List of Document objects with content and metadata
        embeddings: Embedding model
        group_by_source: If True, organize vector store in source subdirectories (default: True)
        
    Returns:
        FAISS vector store
    """
    if group_by_source and chunks:
        # Determine source group from chunks
        source_group_key = _get_source_group_key(chunks)
        source_dir = get_or_create_source_vector_db_dir(source_group_key)
        vector_db_path = source_dir / "store.faiss"
    else:
        # Legacy: store at top level
        vector_db_path = Path(VECTOR_DB_DIR) / f"{filename}.faiss"

    if vector_db_path.exists():
        vector_store = FAISS.load_local(str(vector_db_path), embeddings=embeddings, allow_dangerous_deserialization=True)
    else:
        single_vector = embeddings.embed_query("initialize")
        index = faiss.IndexFlatL2(len(single_vector))
        vector_store = FAISS(
            embedding_function=embeddings,
            index=index,
            docstore=InMemoryDocstore(),
            index_to_docstore_id={}
        )

        try:
            # Add documents - handle both Document objects and strings
            docs_to_add = []
            for chunk in chunks:
                if isinstance(chunk, Document):
                    docs_to_add.append(chunk)
                else:
                    # Convert string to Document with empty metadata
                    docs_to_add.append(Document(page_content=str(chunk), metadata={}))
            
            vector_store.add_documents(docs_to_add)
        except Exception as e:
            print(e)
        else:
            vector_store.save_local(str(vector_db_path))
    return vector_store


def create_retriever_with_filter(vector_store, filter_dict: dict = None, k: int = 4):
    """Create a retriever with optional metadata filtering.
    
    Args:
        vector_store: FAISS vector store
        filter_dict: Dictionary of metadata filters (e.g., {"source_type": "csv", "filename": "data.csv"})
        k: Number of documents to retrieve
        
    Returns:
        Retriever with metadata filtering
    """
    if filter_dict:
        return vector_store.as_retriever(
            search_kwargs={
                "k": k,
                "filter": filter_dict
            }
        )
    else:
        return vector_store.as_retriever(search_kwargs={"k": k})


def filter_documents_by_metadata(documents: list, metadata_filters: dict) -> list:
    """Filter documents by metadata criteria.
    
    Args:
        documents: List of Document objects
        metadata_filters: Dictionary with metadata filter criteria
                         Example: {"source_type": "csv", "filename": "data.csv"}
        
    Returns:
        Filtered list of documents
    """
    filtered = []
    for doc in documents:
        match = True
        for key, value in metadata_filters.items():
            if key not in doc.metadata:
                match = False
                break
            # Support both exact match and substring match
            if isinstance(value, str):
                if value.lower() not in str(doc.metadata[key]).lower():
                    match = False
                    break
            else:
                if doc.metadata[key] != value:
                    match = False
                    break
        
        if match:
            filtered.append(doc)
    
    return filtered


# Build RAG chain
def build_rag_chain(
    retriever,
    model_name=DEFAULT_CHAT_MODEL,
    base_url=OLLAMA_BASE_URL,
    response_type="Plain Text",
):
    prompt = """
        You are an assistant for financial data analysis. Use the retrieved context to answer questions.
        If you don't know the answer, say so.

        Response format requirement: {response_type}
        Formatting rules:
        - Plain Text: respond with plain text only. Do not use markdown lists, headings, or code fences.
        - Markdown: respond using clear markdown formatting.
        - Python Code: respond with Python code only inside a single fenced python code block, with no extra explanation outside the code block.

        Question: {question}
        Context: {context}
        Answer:
    """
    prompt_template = ChatPromptTemplate.from_template(prompt)
    model = ChatOllama(model=model_name, base_url=base_url)
    return (
        {
            "context": retriever | (lambda docs: "\n\n".join(doc.page_content for doc in docs)),
            "question": RunnablePassthrough(),
            "response_type": lambda _: response_type,
        }
        | prompt_template
        | model
        | StrOutputParser()
    )
