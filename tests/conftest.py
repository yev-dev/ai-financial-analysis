"""Pytest configuration and shared fixtures."""
import pytest
import tempfile
from pathlib import Path


@pytest.fixture
def temp_dir():
    """Provide a temporary directory for tests."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


@pytest.fixture
def sample_csv_file(temp_dir):
    """Create a sample CSV file for testing."""
    import pandas as pd
    
    csv_path = temp_dir / "sample_data.csv"
    df = pd.DataFrame({
        "ticker": ["AAPL", "GOOGL", "MSFT"],
        "price": [150.25, 140.50, 380.75],
        "volume": [1000000, 500000, 750000]
    })
    df.to_csv(csv_path, index=False)
    return csv_path


@pytest.fixture
def sample_html_content():
    """Provide sample HTML content for testing."""
    return """
    <html>
        <head>
            <title>Sample Financial Report</title>
        </head>
        <body>
            <article>
                <h1>Financial Analysis</h1>
                <p>This is a sample financial report.</p>
                <h2>Section 1</h2>
                <p>Some financial data here.</p>
            </article>
        </body>
    </html>
    """


@pytest.fixture
def sample_html_file(temp_dir, sample_html_content):
    """Create a sample HTML file for testing."""
    html_path = temp_dir / "sample_report.html"
    with open(html_path, 'w') as f:
        f.write(sample_html_content)
    return html_path


@pytest.fixture
def sample_json_data():
    """Provide sample JSON data for testing."""
    return {
        "company": "Acme Corp",
        "ticker": "ACME",
        "financials": {
            "revenue": 1000000,
            "expenses": 600000,
            "profit": 400000
        },
        "stocks": [
            {"date": "2026-01-01", "price": 100.50, "volume": 1000000},
            {"date": "2026-01-02", "price": 102.25, "volume": 1500000},
            {"date": "2026-01-03", "price": 101.75, "volume": 800000}
        ]
    }


@pytest.fixture
def sample_json_file(temp_dir, sample_json_data):
    """Create a sample JSON file for testing."""
    import json
    
    json_path = temp_dir / "sample_data.json"
    with open(json_path, 'w') as f:
        json.dump(sample_json_data, f, indent=2)
    return json_path


@pytest.fixture
def sample_xml_content():
    """Provide sample XML content for testing."""
    return """
    <report company="Acme Corp" ticker="ACME">
        <financials>
            <revenue currency="USD">1000000</revenue>
            <expenses currency="USD">600000</expenses>
            <profit currency="USD">400000</profit>
        </financials>
        <stocks>
            <stock>
                <date>2026-01-01</date>
                <price>100.50</price>
                <volume>1000000</volume>
            </stock>
            <stock>
                <date>2026-01-02</date>
                <price>102.25</price>
                <volume>1500000</volume>
            </stock>
        </stocks>
    </report>
    """


@pytest.fixture
def sample_xml_file(temp_dir, sample_xml_content):
    """Create a sample XML file for testing."""
    xml_path = temp_dir / "sample_data.xml"
    xml_path.write_text(sample_xml_content, encoding="utf-8")
    return xml_path


@pytest.fixture
def sample_word_file(temp_dir):
    """Create a sample Word file for testing."""
    from docx import Document

    docx_path = temp_dir / "sample_report.docx"
    document = Document()
    document.add_heading("Quarterly Financial Report", level=1)
    document.add_paragraph("Revenue increased year over year.")
    document.add_heading("Metrics", level=2)
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Revenue"
    table.rows[1].cells[1].text = "$1.2M"
    document.save(docx_path)
    return docx_path
